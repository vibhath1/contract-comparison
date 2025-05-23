import os
from pathlib import Path
import PyPDF2
import docx
import pytesseract
from PIL import Image
import pdf2image
import spacy
import uuid
from app.models.schemas import DocumentType

# Initialize spaCy - you might need to install the model first with:
# python -m spacy download en_core_web_lg
try:
    nlp = spacy.load("en_core_web_lg")
except OSError:
    # Fallback to a smaller model if the large one is not available
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # Simplified NLP for basic functionality if spaCy models aren't available
        nlp = None

class DocumentProcessor:
    def __init__(self, upload_dir: Path):
        self.upload_dir = upload_dir
    
    def process_document(self, file_path: Path):
        """Process a document and extract its content."""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._process_docx(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return self._process_image(file_path)
        elif file_extension in ['.txt']:
            return self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: Path):
        """Extract text from PDF files. Uses OCR if needed."""
        try:
            # Try to extract text directly
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text() or ""
            
            # If text extraction returns very little content, try OCR
            if len(text.strip()) < 100:
                return self._ocr_pdf(file_path)
            
            return {
                "text": text,
                "type": DocumentType.PDF,
                "pages": len(pdf_reader.pages)
            }
        except Exception as e:
            # Fallback to OCR on exception
            print(f"Error extracting text from PDF: {str(e)}. Falling back to OCR.")
            return self._ocr_pdf(file_path)
    
    def _ocr_pdf(self, file_path: Path):
        """Process PDF using OCR."""
        try:
            images = pdf2image.convert_from_path(file_path)
            text = ""
            for image in images:
                text += pytesseract.image_to_string(image)
            
            return {
                "text": text,
                "type": DocumentType.PDF,
                "pages": len(images),
                "processed_with_ocr": True
            }
        except Exception as e:
            # If OCR fails, return an error message in the text
            print(f"Error processing PDF with OCR: {str(e)}")
            return {
                "text": f"Error processing document: {str(e)}",
                "type": DocumentType.PDF,
                "pages": 0,
                "error": str(e)
            }
    
    def _process_docx(self, file_path: Path):
        """Extract text and structure from DOCX files."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "text": text,
                "type": DocumentType.DOCX,
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            print(f"Error processing DOCX: {str(e)}")
            return {
                "text": f"Error processing document: {str(e)}",
                "type": DocumentType.DOCX,
                "paragraphs": 0,
                "error": str(e)
            }
    
    def _process_image(self, file_path: Path):
        """Process images using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text,
                "type": DocumentType.IMAGE,
                "processed_with_ocr": True
            }
        except Exception as e:
            print(f"Error processing image: {str(e)}")
            return {
                "text": f"Error processing document: {str(e)}",
                "type": DocumentType.IMAGE,
                "error": str(e)
            }
    
    def _process_text(self, file_path: Path):
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                text = file.read()
            
            return {
                "text": text,
                "type": DocumentType.TEXT
            }
        except Exception as e:
            print(f"Error processing text file: {str(e)}")
            return {
                "text": f"Error processing document: {str(e)}",
                "type": DocumentType.TEXT,
                "error": str(e)
            }
    
    def extract_entities(self, text: str):
        """Extract entities from text using spaCy."""
        if nlp is None:
            # Return empty results if spaCy is not available
            return {
                "dates": [],
                "monetary_values": [],
                "organizations": [],
                "persons": [],
                "locations": []
            }
            
        doc = nlp(text)
        entities = {
            "dates": [],
            "monetary_values": [],
            "organizations": [],
            "persons": [],
            "locations": []
        }
        
        for ent in doc.ents:
            if ent.label_ == "DATE":
                entities["dates"].append(ent.text)
            elif ent.label_ == "MONEY":
                entities["monetary_values"].append(ent.text)
            elif ent.label_ == "ORG":
                entities["organizations"].append(ent.text)
            elif ent.label_ == "PERSON":
                entities["persons"].append(ent.text)
            elif ent.label_ == "GPE":
                entities["locations"].append(ent.text)
        
        return entities